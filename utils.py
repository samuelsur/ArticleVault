import os
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from urllib.parse import urlparse
import mimetypes # mime types is used to determine the type of file being downloaded
from typing import List, Optional
from pydantic import BaseModel, Field
from PIL import Image
from logging_config import setup_logger, setup_file_logger

# Initialize logger
logger = setup_logger(__name__)
file_logger = setup_file_logger(__name__, "logs/article_extractor.log")



class MSWord:

    @staticmethod
    def _get_callback(callback=None):
        """
        Returns the provided callback or a no-op function if None is provided.
        This simplifies progress reporting by eliminating repeated null checks.
        
        Args:
            callback (callable, optional): The callback function to use
        Returns:
            callable: Either the provided callback or a no-op function
        """
        if callback is None:
            return lambda message, percent: None  # No-op function
        return callback
    
    @staticmethod
    def _add_page_border(doc, border_color=(150, 42, 46), border_width=200):
        """
        Add a page border to the document.
        
        Args:
            doc: Document object
            border_color (tuple): RGB color tuple (default: brown)
            border_width (int): Border width in points
        """
        # Convert RGB to hex color
        r, g, b = border_color
        hex_color = f"{r:02x}{g:02x}{b:02x}"
        
        # Add page borders to each section
        for section in doc.sections:
            sectPr = section._sectPr
            
            # Create page borders element
            pgBorders = OxmlElement('w:pgBorders')
            pgBorders.set(qn('w:offsetFrom'), 'page')
            
            # Add each border (top, right, bottom, left)
            for border_type in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_type}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), str(border_width))
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), hex_color)
                pgBorders.append(border)
            
            # Add the borders to the section properties
            if sectPr.find(qn('w:pgBorders')) is not None:
                sectPr.remove(sectPr.find(qn('w:pgBorders')))
            sectPr.append(pgBorders)
    

    @staticmethod
    def create_page_bordered_docx(uuid, filename, content, url, border_color=(150, 42, 46), 
                                border_width=200, header=None, footer=None, 
                                progress_callback=None, user_context=None):
        """
        Create a DOCX file with a border around the entire page with optional header and footer.
        
        Args:
            uuid (str): Unique identifier for the extraction run
            filename (str): Path where to save the document
            content (dict): Document content structure
            url (str): Source URL of the article
            border_color (tuple): RGB color tuple for border (default: brown)
            border_width (int): Border width in points
            header (str, optional): Optional header text
            footer (str, optional): Optional footer text
            progress_callback (callable, optional): Callback function to report progress
            user_context (str, optional): User context for logging purposes
        """
        progress_callback = MSWord._get_callback(progress_callback)
        context_info = f" for {user_context}" if user_context else ""
        
        # temp/uuid/temp_images
        image_session_path = f"temp/{uuid}/temp_images"
        docx_session_path = f"temp/{uuid}/{filename}"
        if not os.path.exists(os.path.dirname(f'temp/{uuid}')):
            os.makedirs(os.path.dirname(f'temp/{uuid}'))
        # Download images if provided
        title = content.get("title")
        image_urls = []
        for block in content["content_blocks"]:
            if block["type"] == "image":
                image_urls.append(block)

        # download the images
        if image_urls:
            progress_callback("Downloading images...", 65)
            images = MSWord._temp_download_images(images=image_urls, download_path=image_session_path, 
                                                progress_callback=progress_callback, user_context=user_context)
            progress_callback("Images downloaded successfully", 75)
        else:
            images = {}

        # Create a new Document
        progress_callback("Creating document structure...", 80)
            
        doc = Document()
        
        # Set page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add header if provided
        if header:
            header_section = doc.sections[0]
            header_obj = header_section.header
            header_para = header_obj.paragraphs[0]
            header_para.text = header
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # date (if available)
        date = content.get("date", "")
        if date:
            date_paragraph = doc.add_paragraph()
            date_run = date_paragraph.add_run(date)
            date_run.italic = True
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_paragraph.space_after = Pt(6)


        # url
        url_paragraph = doc.add_paragraph()
        url_run = url_paragraph.add_run(url)
        url_run.italic = True
        url_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        url_paragraph.space_after = Pt(12)  # Add space after the URL paragraph
        
        
        # Process content blocks
        total_blocks = len(content["content_blocks"])
        for i, block in enumerate(content["content_blocks"]):
            if i % 5 == 0:  # Update progress every 5 blocks
                progress_pct = 80 + (i / total_blocks * 10)  # Progress from 80% to 90%
                progress_callback(f"Creating document content ({i}/{total_blocks})...", progress_pct)
    
            block_type = block["type"]
            if block_type == "text":
                text_content = block["content"]
                if text_content:
                    # Create a new paragraph for the text
                    paragraph = doc.add_paragraph()
                    
                    # Check for markdown-style formatting indicators
                    
                    # Handle headings (lines starting with ## or #)
                    if text_content.lstrip().startswith('##'):
                        run = paragraph.add_run(text_content.lstrip('#').strip())
                        run.bold = True
                        run.font.size = Pt(14)
                        paragraph.style = 'Heading 2'
                    elif text_content.lstrip().startswith('#'):
                        run = paragraph.add_run(text_content.lstrip('#').strip())
                        run.bold = True
                        run.font.size = Pt(15)
                        paragraph.style = 'Heading 1'
                    
                    # Handle entire paragraph in bold (**text**)
                    elif text_content.startswith('**') and text_content.endswith('**'):
                        run = paragraph.add_run(text_content.strip('**'))
                        run.bold = True
                    
                    # Handle entire paragraph in italics (_text_)
                    elif (text_content.startswith('_') and text_content.endswith('_')) or \
                         (text_content.startswith('*') and text_content.endswith('*')):
                        run = paragraph.add_run(text_content.strip('_*'))
                        run.italic = True
                    
                    # Handle mixed formatting with ** for bold and _ or * for italic
                    else:
                        # Process text with inline formatting
                        remaining_text = text_content
                        
                        # Iterator for inline formatting
                        while remaining_text:
                            # Look for bold text
                            bold_start = remaining_text.find('**')
                            if bold_start != -1:
                                bold_end = remaining_text.find('**', bold_start + 2)
                                if bold_end != -1:
                                    # Add text before the bold section
                                    if bold_start > 0:
                                        run = paragraph.add_run(remaining_text[:bold_start])
                                    
                                    # Add the bold text
                                    bold_text = remaining_text[bold_start + 2:bold_end]
                                    run = paragraph.add_run(bold_text)
                                    run.bold = True
                                    
                                    # Continue with remaining text
                                    remaining_text = remaining_text[bold_end + 2:]
                                    continue
                            
                            # Look for italic text
                            italic_marker = '_' if '_' in remaining_text else '*'
                            italic_start = remaining_text.find(italic_marker)
                            if italic_start != -1:
                                italic_end = remaining_text.find(italic_marker, italic_start + 1)
                                if italic_end != -1:
                                    # Add text before the italic section
                                    if italic_start > 0:
                                        run = paragraph.add_run(remaining_text[:italic_start])
                                    
                                    # Add the italic text
                                    italic_text = remaining_text[italic_start + 1:italic_end]
                                    run = paragraph.add_run(italic_text)
                                    run.italic = True
                                    
                                    # Continue with remaining text
                                    remaining_text = remaining_text[italic_end + 1:]
                                    continue
                            
                            # No more formatting found, add the rest of the text
                            if remaining_text:
                                paragraph.add_run(remaining_text)
                                break

            elif block_type == "image":
                url = block["url"]
                caption = block.get("caption", "")
                # get image path from the downloaded images
                image_path = images.get(url, None)
                
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if image_path and os.path.exists(image_path):
                    try:
                        run = p.add_run()
                        run.add_picture(image_path, width=Inches(4)) # the image width is hardcoded to 4 inches, todo: make it configurable in streamlit
                        
                        # Add caption if available
                        caption = block.get('caption', '')
                        if caption:
                            caption_para = doc.add_paragraph()
                            caption_run = caption_para.add_run(caption)
                            caption_run.italic = True
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    except Exception as img_error:
                        # If adding the image fails, add a placeholder
                        error_msg = f"Error adding image to document{context_info}: {url} - {str(img_error)}"
                        logger.error(error_msg)
                        file_logger.error(error_msg)
                        run = p.add_run("[Image could not be displayed]")
                        run.italic = True
                        
                        # Try to still add caption
                        if caption:
                            p.add_run("\n" + caption)
                else:
                    # Add a placeholder if image wasn't downloaded successfully
                    run = p.add_run("[Image unavailable]")
                    run.italic = True

        if image_urls:
            # Clean up temporary images after adding them to the document
            MSWord._temp_delete_images(download_path=image_session_path, user_context=user_context)
        
        # Add footer if provided
        if footer:
            footer_section = doc.sections[0]
            footer_obj = footer_section.footer
            footer_para = footer_obj.paragraphs[0]
            footer_para.text = footer
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page border to the document
        MSWord._add_page_border(doc, border_color, border_width)
        
        progress_callback("Saving document...", 95)
            
        # Save the document
        doc.save(docx_session_path)
        
        progress_callback("Document saved successfully", 98)
            
        return f"Document with page border saved as {filename}"


    def _temp_delete_images(download_path: str = "temp_images", user_context: str = None):
        """
        Delete temporary images downloaded during the document creation.
        This method should be called after the document is saved to clean up.
        
        Args:
            download_path (str): Path to the temporary images directory
            user_context (str, optional): User context for logging purposes
        """
        import os
        import shutil
        
        context_info = f" for {user_context}" if user_context else ""
        
        if os.path.exists(download_path):
            shutil.rmtree(download_path)
            info_msg = f"Temporary images deleted from {download_path}{context_info}"
            logger.info(info_msg)
            file_logger.info(info_msg)
        else:
            info_msg = f"Temporary images directory does not exist, nothing to delete{context_info}"
            logger.info(info_msg)
            file_logger.info(info_msg)
   
    @staticmethod
    def _temp_download_images(images: list[dict], download_path: str = "temp_images", 
                            progress_callback=None, user_context: str = None) -> dict:
        """
        Download images from the provided URLs and save them to the specified path.
        
        Args:
            images (list[dict]): List of dictionaries containing image URLs and metadata.
            download_path (str): Path where to save the downloaded images
            progress_callback (callable, optional): Callback function to report progress
            user_context (str, optional): User context for logging purposes
        Returns:
            dict: A dictionary mapping image URLs to their local file paths
        """
        progress_callback = MSWord._get_callback(progress_callback)
        context_info = f" for {user_context}" if user_context else ""
        
        
        if not os.path.exists(download_path):
            os.makedirs(download_path)
        
        # Add Pillow import for image conversion
        try:
            from PIL import Image
            import io
            has_pillow = True
        except ImportError:
            warning_msg = f"Pillow library not found{context_info}. Image conversion will be limited."
            logger.warning(warning_msg)
            file_logger.warning(warning_msg)
            has_pillow = False
        
        im_dict = dict()
        total_images = len(images)
        
        for i, image in enumerate(images):
            img_url = image["url"]
            
            if total_images > 0:
                progress_pct = 65 + ((i / total_images) * 10)  # Progress from 65% to 75%
                progress_callback(f"Downloading image {i+1}/{total_images}", progress_pct)
                
            try:
                parsed_url = urlparse(img_url)
                file_name = os.path.basename(parsed_url.path)
                response = requests.get(img_url)
                response.raise_for_status()
                
                # Generate a unique filename to avoid collisions
                base_name = os.path.splitext(file_name)[0]
                safe_name = f"{base_name}_{hash(img_url) % 10000}"
                
                # Path for the original download
                temp_path = os.path.join(download_path, f"temp_{safe_name}")
                
                # Save the original image data first
                with open(temp_path, 'wb') as img_file:
                    img_file.write(response.content)
                
                if has_pillow:
                    try:
                        # Try to open and convert the image with Pillow
                        img = Image.open(temp_path)
                        
                        # Convert to RGB mode (handles RGBA, grayscale, etc.)
                        if img.mode != 'RGB':
                            img = img.convert('RGB')
                        
                        # Save as PNG (format python-docx handles well)
                        final_path = os.path.join(download_path, f"{safe_name}.png")
                        img.save(final_path, "PNG")
                        
                        # Clean up the temp file
                        os.remove(temp_path)
                        
                        im_dict[img_url] = final_path
                        info_msg = f"Image {img_url} converted and saved as {final_path}{context_info}"
                        logger.info(info_msg)
                        file_logger.info(info_msg)
                        
                    except Exception as img_error:
                        error_msg = f"Image conversion failed{context_info} for {img_url}: {img_error}"
                        logger.error(error_msg)
                        file_logger.error(error_msg)
                        # If conversion fails, try using the original file
                        content_type = response.headers.get('Content-Type', '')
                        extension = mimetypes.guess_extension(content_type) or '.jpg'
                        final_path = os.path.join(download_path, f"{safe_name}{extension}")
                        os.rename(temp_path, final_path)
                        im_dict[img_url] = final_path
                else:
                    # Without Pillow, just guess the extension and hope for the best
                    content_type = response.headers.get('Content-Type', '')
                    extension = mimetypes.guess_extension(content_type) or '.jpg'
                    final_path = os.path.join(download_path, f"{safe_name}{extension}")
                    os.rename(temp_path, final_path)
                    im_dict[img_url] = final_path
                    
            except Exception as e:
                error_msg = f"Failed to download {img_url}{context_info}: {e}"
                logger.error(error_msg)
                file_logger.error(error_msg)


        return im_dict  # Return the dictionary of downloaded images with their paths




class ContentBlock(BaseModel):
    type: str = Field(description="Type of content block: 'text' or 'image'")
    content: Optional[str] = Field(None, description="Text content (if type is 'text')")
    url: Optional[str] = Field(None, description="Image URL (if type is 'image')")
    alt_text: Optional[str] = Field(None, description="Alternative text for image")
    caption: Optional[str] = Field(None, description="Caption for image")

class Article(BaseModel):
    title: str = Field(description="Title of the extracted article")
    date: str = Field(description="Publication date of the article (if available)")
    author: str = Field(description="Author of the article (if available)")
    content_blocks: List[ContentBlock] = Field(
        default_factory=list,
        description="List of content blocks (text or images) in order of appearance"
    )


if __name__ == "__main__":
    pass
